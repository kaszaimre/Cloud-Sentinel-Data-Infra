"# ==============================================================================\n"
        module_desc = """ 
        "# ARCHÍVUM NÉV: SZOVEGARHIV_MASTER_LOG.docx\n"
        f"# DÁTUM: {datetime.now().strftime('%Y-%m-%d')}\n"
        f"# FELDOLGOZOTT MENNYISÉG: {total_images} darab fotó / images\n"
        "#\n"

        "# LEÍRÁS ÉS FELADAT / DESCRIPTION AND TASK:\n"
        "# A 'SZOVEGARHIV' digitális gyűjtemény automatizált, rendszerszintű feldolgozása.\n"
        "# Megakadályozza az információvesztést azáltal, hogy az 5000+ képernyőfotóból és\n"
        "# dokumentum-kivonatból álló 'Borsodi' master stratégia szöveges anyagát\n"
        "# optikai karakterfelismeréssel (OCR) egyetlen kereshető Word dokumentummá fűzi.\n"
        "#\n"
        "# This module serves as the automated compilation of the 'SZOVEGARHIV' repository.\n"
        "# It prevents information loss by consolidating the textual data of the 'Borsodi'\n"
        "# master strategy from 5000+ screenshots into a single searchable Word file\n"
        "# utilizing Optical Character Recognition (OCR) engines.\n"
        SZERZŐ: Don Mérnök (Tábornok) | BORSODI WAR ROOM """
"# ==============================================================================\n"

import os
import glob
from datetime import datetime
from PIL import Image
import pytesseract
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Tesseract elérési útja
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def add_header_description(doc):
    """Létrehozza a professzionális magyar és angol nyelvű fejléc leírást"""
    # Főcím
    title = doc.add_heading('SZOVEGARHIV', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metaadatok táblázat vagy blokk formában
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run(f"Generálva / Generated: {datetime.now().strftime('%Y-%m-%d')}\nForrás / Source: Google Fotók - SZOVEGARHIV Album\n")
    run_meta.italic = True
    run_meta.font.size = Pt(10)
    
    # --- MAGYAR LEÍRÁS ---
    p_hu = doc.add_paragraph()
    p_hu.paragraph_format.space_before = Pt(12)
    p_hu.paragraph_format.space_after = Pt(6)
    
    run_hu_title = p_hu.add_run("Dokumentum Leírás (HU):\n")
    run_hu_title.bold = True
    run_hu_title.font.color.rgb = RGBColor(0, 51, 102) # Sötétkék kiemelés
    
    run_hu_text = p_hu.add_run(
        "Ez a dokumentum a 'SZOVEGARHIV' elnevezésű digitális képgyűjtemény automatizált, "
        "szövegesített változata. A tartalom 5000+ képernyőfotóból és dokumentum-kivonatból került "
        "összeállításra optikai karakterfelismerő (OCR) technológia segítségével. "
        "A szövegek hűen tükrözik a 'Borsodi' master stratégia, a rendszerszintű figyelemirányítás "
        "és a kiber-térbeli műveletek operatív tudatosságát. Az eredeti forrásfájlok nevei az egyes "
        "fejezeteknél hivatkozásként szerepelnek a visszakereshetőség érdekében."
    )
    run_hu_text.font.size = Pt(11)
    
    # --- ANGOL LEÍRÁS ---
    p_en = doc.add_paragraph()
    p_en.paragraph_format.space_before = Pt(6)
    p_en.paragraph_format.space_after = Pt(24)
    
    run_en_title = p_en.add_run("Document Description (EN):\n")
    run_en_title.bold = True
    run_en_title.font.color.rgb = RGBColor(0, 51, 102)
    
    run_en_text = p_en.add_run(
        "This document is the automated, text-based compilation of the digital image collection titled 'SZOVEGARHIV'. "
        "The content has been extracted from 5000+ screenshots and document excerpts using Optical Character Recognition (OCR) technology. "
        "The texts faithfully log the operational awareness of the 'Borsodi' master strategy, system-level attention management, "
        "and cyber-space operations. For traceability, the original source filenames are referenced at the beginning of each section."
    )
    run_en_text.font.size = Pt(11)
    
    # Elválasztó vonal imitáció (stílusos lezárás a tartalom előtt)
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep.add_run("=" * 60).font.color.rgb = RGBColor(128, 128, 128)

def tomeges_ocr_wordbe(kepek_mappaja, kimeneti_word_fajl):
    doc = Document()
    
    # Fejléc hozzáadása a dokumentum elejére
    add_header_description(doc)
    
    # Képek keresése
    kep_formatumok = ["*.jpg", "*.jpeg", "*.png", "*.JPG", "*.JPEG", "*.PNG"]
    fajlok = []
    for ext in kep_formatumok:
        fajlok.extend(glob.glob(os.path.join(kepek_mappaja, ext)))
    
    fajlok.sort()
    osszes_kep = len(fajlok)
    print(f"Összesen {osszes_kep} képet találtam. Indul a pörkölés... 🔥")
    
    for index, kep_utvonal in enumerate(fajlok, start=1):
        fajl_neve = os.path.basename(kep_utvonal)
        print(f"[{index}/{osszes_kep}] Feldolgozás: {fajl_neve}...")
        
        try:
            kep = Image.open(kep_utvonal)
            szoveg = pytesseract.image_to_string(kep, lang='hun')
            
            # Cím a kép nevével
            h = doc.add_heading(f"Forrás kép / Source file: {fajl_neve}", level=2)
            h.paragraph_format.space_before = Pt(18)
            
            # Szövegtörzs beillesztése
            doc.add_paragraph(szoveg)
            doc.add_page_break()
            
        except Exception as e:
            print(f"⚠️ Hiba a {fajl_neve} fájlnál: {e}")
            doc.add_heading(f"HIBA / ERROR: {fajl_neve}", level=2)
            doc.add_paragraph(f"Sikertelen OCR művelet. / OCR failed. Hiba: {e}")
            doc.add_page_break()

    doc.save(kimeneti_word_fajl)
    print(f"\n✅ KÉSZ! Mentve ide: {kimeneti_word_fajl}")

"# ==============================================================================\n"
"#def hosszú_kép_feldolgozás(kep_utvonal, max_magassag=2000):
"#   """Felszeleteli a túl hosszú kollázsokat a pontosabb szövegfelismerésért"""
"#    from PIL import Image
"#   import pytesseract
"#    
"#    teljes_szoveg = ""
    with Image.open(kep_utvonal) as img:
"#        szelesseg, magassag = img.size
        
        # Ha a kép hosszabb, mint a biztonságos limit, daraboljuk
        if magassag > max_magassag:
            szeletek = magassag // max_magassag + 1
            for i in range(szeletek):
                felso = i * max_magassag
                also = min((i + 1) * max_magassag, magassag)
                
                # Szelet kivágása a memóriában
                szelet = img.crop((0, felso, szelesseg, also))
                szoveg_resz = pytesseract.image_to_string(szelet, lang='hun')
                teljes_szoveg += szoveg_resz + "\n"
        else:
            teljes_szoveg = pytesseract.image_to_string(img, lang='hun')
            
    return teljes_szoveg

# Beállítások
KEPEK_DIR = r"C:\kepek" 
KIMENET_DOCX = r"C:\kepek\teljes_szovegarhiv.docx"

if __name__ == "__main__":
    tomeges_ocr_wordbe(KEPEK_DIR, KIMENET_DOCX)
    
    Ha ezt beépíted az ocr_to_word_enterprise.py fő ciklusába a sima image_to_string helyére, akkor az 
    összes ilyen monumentális képernyőkép-sorozatot másodpercek alatt át tudod nyomni a Word dokumentumba
"# ==============================================================================\n"
